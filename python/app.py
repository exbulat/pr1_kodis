import os
import io
import json
import uuid
import warnings
import logging

warnings.filterwarnings("ignore")
logging.getLogger("transformers").setLevel(logging.ERROR)

from datetime import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

from transcriber import transcribe_audio
from analyzer import analyze_transcript

load_dotenv()

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "..", "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXTS = {".mp3", ".wav", ".mp4", ".m4a", ".ogg", ".webm", ".avi", ".mov"}


def _clean(name: str) -> str:
    keep = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789_-.")
    return "".join(c if c in keep else "_" for c in name)


def _build_docx(data: dict) -> io.BytesIO:
    doc = Document()
    meeting = data.get("meeting", {})
    title = meeting.get("title") or "Протокол совещания"
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")

    if data.get("sentiment"):
        p = doc.add_paragraph(f"Тон: {data['sentiment']}")
        p.runs[0].italic = True
    if data.get("next_meeting"):
        doc.add_paragraph(f"Следующее совещание: {data['next_meeting']}")

    participants = (meeting.get("participants") or [])
    if participants:
        doc.add_heading("Участники", level=1)
        for p in participants:
            name = p.get("name") or "—"
            role = p.get("role") or ""
            doc.add_paragraph(f"{name} — {role}" if role else name, style="List Bullet")

    if data.get("summary"):
        doc.add_heading("Краткое содержание", level=1)
        doc.add_paragraph(data["summary"])

    kw = data.get("keywords") or []
    if kw:
        doc.add_heading("Ключевые темы", level=1)
        doc.add_paragraph(", ".join(kw))

    contacts = data.get("contacts") or []
    if contacts:
        doc.add_heading("Контакты", level=1)
        for c in contacts:
            parts = [c.get("name") or ""]
            for f in ("position", "email", "phone"):
                if c.get(f): parts.append(c[f])
            doc.add_paragraph(" · ".join(filter(None, parts)), style="List Bullet")

    decisions = data.get("decisions") or []
    if decisions:
        doc.add_heading("Принятые решения", level=1)
        for i, d in enumerate(decisions, 1):
            doc.add_paragraph(f"{i}. {d}")

    tasks = data.get("tasks") or []
    if tasks:
        doc.add_heading("Задачи", level=1)
        for t in tasks:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(t.get("task") or "").bold = True
            p.add_run(f"\n    {t.get('assignee','—')} · срок: {t.get('deadline','не указан')} · приоритет: {t.get('priority','—')}")

    if data.get("transcript"):
        doc.add_heading("Транскрипт", level=1)
        for chunk in data["transcript"].split("\n"):
            if chunk.strip():
                doc.add_paragraph(chunk.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/health")
def health():
    return jsonify({"status": "ok"})


@app.route("/analyze_text", methods=["POST"])
def analyze_text():
    data = request.get_json()
    text = (data or {}).get("text", "").strip()
    if not text:
        return jsonify({"error": "Текст пустой"}), 400
    try:
        result = analyze_transcript(text)
        result["transcript"] = text
        result["has_diarization"] = False
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/analyze_audio", methods=["POST"])
def analyze_audio():
    if "file" not in request.files:
        return jsonify({"error": "Файл не найден"}), 400

    file = request.files["file"]
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTS:
        return jsonify({"error": f"Формат {ext} не поддерживается"}), 400

    safe_name = f"{uuid.uuid4().hex}_{_clean(file.filename)}"
    file_path = os.path.join(UPLOAD_FOLDER, safe_name)
    file.save(file_path)

    try:
        transcription = transcribe_audio(file_path)
        segments = transcription.get("segments", [])
        text = " ".join(s["text"] for s in segments) if segments else transcription.get("text", "")

        result = analyze_transcript(text)
        result["transcript"] = text
        result["has_diarization"] = False
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try: os.remove(file_path)
        except: pass


@app.route("/export_word", methods=["POST"])
def export_word():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Нет данных"}), 400
    try:
        buf = _build_docx(data)
        title = (data.get("meeting") or {}).get("title") or "protocol"
        filename = f"{_clean(title)[:40]}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(buf, as_attachment=True, download_name=filename,
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    print("Сервер запущен: http://127.0.0.1:5000")
    app.run(host="127.0.0.1", port=5000, debug=False, use_reloader=False)
